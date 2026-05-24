import sys
import os
import argparse
import uuid
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_args():
    parser = argparse.ArgumentParser(description="PDF to DOCX Converter")
    parser.add_argument("--input", required=True, help="Path to input PDF file")
    parser.add_argument("--output", required=True, help="Path to output DOCX file")
    parser.add_argument("--mode", default="high", choices=["fast", "high"], help="Conversion mode")
    return parser.parse_args()

def map_font_name(pdf_font):
    pdf_font = pdf_font.lower()
    if "times" in pdf_font or "roman" in pdf_font:
        return "Times New Roman"
    elif "courier" in pdf_font or "mono" in pdf_font:
        return "Courier New"
    elif "arial" in pdf_font:
        return "Arial"
    elif "calibri" in pdf_font:
        return "Calibri"
    elif "helvetica" in pdf_font:
        return "Arial"
    return "Calibri"  # Default fallback

def group_lines_by_y(lines, threshold=6.0):
    rows = []
    for line in sorted(lines, key=lambda l: l["bbox"][1]):
        added = False
        for row in rows:
            # calculate average y-bounds of the current row
            row_y0 = sum(l["bbox"][1] for l in row) / len(row)
            row_y1 = sum(l["bbox"][3] for l in row) / len(row)
            line_y0 = line["bbox"][1]
            line_y1 = line["bbox"][3]
            
            # overlap check (either close y0, or significant height overlap)
            if abs(line_y0 - row_y0) < threshold or (line_y0 < row_y1 and line_y1 > row_y0):
                row.append(line)
                added = True
                break
        if not added:
            rows.append([line])
    return rows

def convert_pdf_to_docx(pdf_path, docx_path, mode="high"):
    print(f"Opening PDF: {pdf_path}")
    doc = fitz.open(pdf_path)
    word_doc = Document()
    
    # Track metrics for admin dashboard integration
    total_pages = len(doc)
    print(f"Total pages: {total_pages}")
    
    # Process page by page
    for page_num in range(total_pages):
        page = doc[page_num]
        rect = page.rect
        page_width_pt = rect.width
        page_height_pt = rect.height
        
        # Add page break for subsequent pages
        if page_num > 0:
            word_doc.add_page_break()
            
        # Get active section for page setup
        section = word_doc.sections[-1]
        section.page_width = Inches(page_width_pt / 72.0)
        section.page_height = Inches(page_height_pt / 72.0)
        
        # Default margins
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Extract text blocks & layout using dict representation
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        
        # Sort blocks top-to-bottom
        blocks = sorted(blocks, key=lambda b: b.get("bbox", [0, 0, 0, 0])[1])
        
        for block in blocks:
            bbox = block.get("bbox", [0, 0, 0, 0])
            block_type = block.get("type", 0)
            
            if block_type == 0:  # Text Block
                lines = block.get("lines", [])
                
                # In high-accuracy mode, we group lines horizontally to detect tables/columns
                if mode == "high" and len(lines) > 1:
                    rows = group_lines_by_y(lines)
                else:
                    rows = [[line] for line in lines]
                    
                for row_lines in rows:
                    # Sort lines horizontally inside the row
                    row_lines = sorted(row_lines, key=lambda l: l["bbox"][0])
                    
                    # Determine if it's a multi-column row (table)
                    is_table = False
                    if len(row_lines) > 1:
                        # Check for horizontal separation gap
                        for i in range(len(row_lines) - 1):
                            if row_lines[i]["bbox"][2] < row_lines[i+1]["bbox"][0]:
                                is_table = True
                                break
                                
                    if is_table:
                        # Create a single-row table in Word to preserve columns
                        cols_count = len(row_lines)
                        table = word_doc.add_table(rows=1, cols=cols_count)
                        table.autofit = True
                        
                        # Populate table cells
                        for idx, line in enumerate(row_lines):
                            cell = table.cell(0, idx)
                            # Remove default paragraph and add styled content
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.space_before = Pt(0)
                            
                            spans = line.get("spans", [])
                            for span in spans:
                                run = p.add_run(span["text"])
                                run.font.name = map_font_name(span["font"])
                                run.font.size = Pt(span["size"])
                                
                                # Font formatting (bold/italic/color)
                                flags = span["flags"]
                                is_italic = bool(flags & 2)
                                is_bold = bool(flags & 16)
                                if is_bold:
                                    run.bold = True
                                if is_italic:
                                    run.italic = True
                                    
                                # Color mapping
                                color_val = span["color"]
                                r = (color_val >> 16) & 255
                                g = (color_val >> 8) & 255
                                b = color_val & 255
                                run.font.color.rgb = RGBColor(r, g, b)
                    else:
                        # Standard paragraph
                        p = word_doc.add_paragraph()
                        # Narrow layout spacing
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        
                        # Set alignment based on block x position relative to page width
                        line = row_lines[0]
                        line_width = line["bbox"][2] - line["bbox"][0]
                        line_center = line["bbox"][0] + (line_width / 2)
                        page_center = page_width_pt / 2
                        
                        if abs(line_center - page_center) < 15 and line_width < (page_width_pt * 0.7):
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif line["bbox"][0] > (page_width_pt * 0.5):
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                        spans = line.get("spans", [])
                        for span in spans:
                            text = span["text"]
                            run = p.add_run(text)
                            run.font.name = map_font_name(span["font"])
                            run.font.size = Pt(span["size"])
                            
                            # Font formatting
                            flags = span["flags"]
                            is_italic = bool(flags & 2)
                            is_bold = bool(flags & 16)
                            if is_bold:
                                run.bold = True
                            if is_italic:
                                run.italic = True
                                
                            color_val = span["color"]
                            r = (color_val >> 16) & 255
                            g = (color_val >> 8) & 255
                            b = color_val & 255
                            run.font.color.rgb = RGBColor(r, g, b)
                            
            elif block_type == 1:  # Image Block
                image_data = block.get("image")
                if image_data:
                    # Determine dimensions
                    img_width = block.get("width", 100)
                    img_height = block.get("height", 100)
                    ext = block.get("ext", "png")
                    
                    # Create temporary image file
                    temp_img_fd, temp_img_path = tempfile.mkstemp(suffix=f".{ext}")
                    try:
                        with os.fdopen(temp_img_fd, "wb") as temp_file:
                            temp_file.write(image_data)
                        
                        # Add image to DOCX
                        p = word_doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        
                        run = p.add_run()
                        # Calculate layout width in inches (fit within margins)
                        max_layout_width_pt = page_width_pt - 108  # 1.5 inch margins total
                        display_width = min(img_width, max_layout_width_pt)
                        
                        run.add_picture(temp_img_path, width=Inches(display_width / 72.0))
                    except Exception as e:
                        print(f"Error inserting image: {e}")
                    finally:
                        # Clean up temp file
                        if os.path.exists(temp_img_path):
                            os.remove(temp_img_path)
                            
    # Save the output DOCX document
    print(f"Saving DOCX to: {docx_path}")
    word_doc.save(docx_path)
    print("Conversion completed successfully!")

if __name__ == "__main__":
    args = parse_args()
    convert_pdf_to_docx(args.input, args.output, args.mode)
